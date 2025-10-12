"""
File extraction utilities for different formats.
Handles PDF, DOCX, CSV, and TXT files.
"""
import io
import csv
import json
from typing import List, Dict, Any, Union
import logging

logger = logging.getLogger(__name__)

def extract_from_txt(file_content: Union[str, bytes]) -> str:
    """Extract text from TXT files."""
    try:
        if isinstance(file_content, bytes):
            # Try different encodings
            for encoding in ['utf-8', 'utf-16', 'latin-1', 'cp1252']:
                try:
                    return file_content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            # If all encodings fail, use utf-8 with errors ignored
            return file_content.decode('utf-8', errors='ignore')
        return str(file_content)
    except Exception as e:
        logger.error(f"Error extracting text from TXT: {str(e)}")
        raise ValueError(f"Failed to extract text from TXT file: {str(e)}")

def extract_from_csv(file_content: Union[str, bytes]) -> str:
    """Extract text from CSV files."""
    try:
        if isinstance(file_content, bytes):
            file_content = file_content.decode('utf-8')
        
        # Parse CSV and extract text from all cells
        reader = csv.reader(io.StringIO(file_content))
        extracted_text = []
        
        for row_num, row in enumerate(reader):
            if row_num == 0:
                # Check if first row looks like headers
                headers = [cell.strip().lower() for cell in row]
                if any(header in ['headline', 'title', 'subject', 'header'] for header in headers):
                    continue  # Skip header row
            
            # Join non-empty cells in each row
            row_text = ' | '.join(cell.strip() for cell in row if cell.strip())
            if row_text:
                extracted_text.append(row_text)
        
        return '\n\n'.join(extracted_text)
    except Exception as e:
        logger.error(f"Error extracting text from CSV: {str(e)}")
        raise ValueError(f"Failed to extract text from CSV file: {str(e)}")

def extract_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF files using pdfminer."""
    try:
        # Try to import pdfminer.six (fallback implementation)
        try:
            from pdfminer.high_level import extract_text
            return extract_text(io.BytesIO(file_content))
        except ImportError:
            # Fallback: basic PDF text extraction (very limited)
            logger.warning("pdfminer.six not available, using basic PDF extraction")
            return extract_basic_pdf_text(file_content)
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        raise ValueError(f"Failed to extract text from PDF file: {str(e)}")

def extract_basic_pdf_text(file_content: bytes) -> str:
    """Basic PDF text extraction fallback (very limited)."""
    try:
        # Very basic extraction - looks for text between BT/ET markers
        content = file_content.decode('latin-1', errors='ignore')
        text_parts = []
        
        # Find text objects in PDF
        start_markers = content.split('BT\n')
        for part in start_markers[1:]:  # Skip first part (before any BT)
            if 'ET' in part:
                text_part = part.split('ET')[0]
                # Basic cleanup - remove PDF commands and extract strings
                lines = text_part.split('\n')
                for line in lines:
                    line = line.strip()
                    # Look for strings in parentheses or brackets
                    if '(' in line and ')' in line:
                        start = line.find('(')
                        end = line.rfind(')')
                        if start < end:
                            extracted = line[start+1:end]
                            if extracted.strip():
                                text_parts.append(extracted)
        
        result = '\n'.join(text_parts)
        return result if result.strip() else "Could not extract text from PDF. Please install pdfminer.six for better PDF support."
    except Exception as e:
        return "Failed to extract text from PDF file."

def extract_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX files."""
    try:
        # Try to import python-docx
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_content))
            
            # Extract text from all paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            return '\n\n'.join(text_parts)
        except ImportError:
            logger.warning("python-docx not available, using basic DOCX extraction")
            return extract_basic_docx_text(file_content)
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        raise ValueError(f"Failed to extract text from DOCX file: {str(e)}")

def extract_basic_docx_text(file_content: bytes) -> str:
    """Basic DOCX text extraction fallback."""
    try:
        import zipfile
        import xml.etree.ElementTree as ET
        
        # DOCX is a ZIP file containing XML documents
        with zipfile.ZipFile(io.BytesIO(file_content)) as docx_zip:
            # Read the main document content
            try:
                doc_xml = docx_zip.read('word/document.xml')
                root = ET.fromstring(doc_xml)
                
                # Extract text from all text nodes
                text_parts = []
                for text_elem in root.iter():
                    if text_elem.tag.endswith('}t'):  # Text elements
                        if text_elem.text:
                            text_parts.append(text_elem.text)
                
                return ''.join(text_parts)
            except Exception:
                return "Could not extract text from DOCX. Please install python-docx for better DOCX support."
    except Exception as e:
        return "Failed to extract text from DOCX file."

def extract_text_from_file(filename: str, file_content: bytes) -> str:
    """
    Extract text from file based on its extension.
    
    Args:
        filename: The original filename
        file_content: The file content as bytes
        
    Returns:
        Extracted text content
    """
    filename_lower = filename.lower()
    
    try:
        if filename_lower.endswith('.txt'):
            return extract_from_txt(file_content)
        elif filename_lower.endswith('.csv'):
            return extract_from_csv(file_content)
        elif filename_lower.endswith('.pdf'):
            return extract_from_pdf(file_content)
        elif filename_lower.endswith(('.docx', '.doc')):
            return extract_from_docx(file_content)
        else:
            # Try to treat as text file
            return extract_from_txt(file_content)
    except Exception as e:
        logger.error(f"Error extracting text from {filename}: {str(e)}")
        raise ValueError(f"Could not extract text from {filename}: {str(e)}")

def get_supported_extensions() -> List[str]:
    """Get list of supported file extensions."""
    return ['.txt', '.csv', '.pdf', '.docx', '.doc']

def is_supported_file(filename: str) -> bool:
    """Check if file extension is supported."""
    return any(filename.lower().endswith(ext) for ext in get_supported_extensions())

class FileExtractor:
    """File extractor class for API integration."""
    
    def extract_from_file(self, file_content: bytes, filename: str, platform: str = 'facebook') -> List[Dict[str, Any]]:
        """Extract and parse ad copy from file content."""
        try:
            # Extract text from file
            extracted_text = extract_text_from_file(filename, file_content)
            
            # Parse the extracted text using TextParser
            from .text_parser import TextParser
            parser = TextParser()
            ads = parser.parse_text(extracted_text, platform)
            
            return ads
        except Exception as e:
            logger.error(f"Error extracting ads from file {filename}: {str(e)}")
            return []
